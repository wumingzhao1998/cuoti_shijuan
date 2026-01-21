import base64
import html
import io
import json
import os
import random
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

import requests
import streamlit as st
from docx import Document
from docx.shared import Inches
from streamlit.runtime.secrets import StreamlitSecretNotFoundError

# 版本信息
VERSION = "1.0"

# 飞书多维表格配置（支持环境变量覆盖）
APP_TOKEN = os.getenv("FEISHU_APP_TOKEN", "NO9nbcpjraKeUCsSQkBcHL9gnhh")
TABLE_ID = os.getenv("FEISHU_TABLE_ID", "tblchSd315sqHTCt")


@st.cache_data(show_spinner=False, ttl=50 * 60)
def get_tenant_access_token(app_id: str, app_secret: str) -> str:
    """
    获取 tenant_access_token，用于后续调用多维表格接口。
    """
    url = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal/"
    resp = requests.post(url, json={"app_id": app_id, "app_secret": app_secret}, timeout=10)
    resp.raise_for_status()
    data = resp.json()
    if data.get("code") != 0:
        raise RuntimeError(f"获取 tenant_access_token 失败: {data}")
    return data["tenant_access_token"]


def fetch_records(token: str) -> List[Dict]:
    """
    拉取表格全部记录，自动翻页。
    """
    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{TABLE_ID}/records/search"
    headers = {"Authorization": f"Bearer {token}"}
    page_token = None
    records: List[Dict] = []

    while True:
        payload: Dict[str, object] = {"page_size": 100}
        if page_token:
            payload["page_token"] = page_token

        resp = requests.post(url, headers=headers, json=payload, timeout=10)
        if not resp.ok:
            # 返回更友好的错误信息，便于排查 token/table 权限问题
            try:
                detail = resp.json()
            except Exception:
                detail = resp.text
            raise RuntimeError(f"拉取记录失败 HTTP {resp.status_code}: {detail}")
        data = resp.json()
        if data.get("code") != 0:
            raise RuntimeError(f"拉取记录失败: {data}")

        items = data["data"].get("items", [])
        records.extend(items)
        page_token = data["data"].get("page_token")

        if not data["data"].get("has_more"):
            break

    return records


def normalize_to_list(value) -> List[str]:
    """
    将单值或列表字段统一转成字符串列表。
    """
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v) for v in value if v is not None]
    return [str(value)]


def normalize_text(value) -> str:
    """
    将字段安全转换为字符串；列表会合并成换行分隔。
    """
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join([str(v) for v in value if v is not None]).strip()
    return str(value).strip()


def is_image_file(name: str, mime: str = None) -> bool:
    """
    判断是否为图片文件（根据MIME类型或文件扩展名）。
    """
    if mime:
        return mime.startswith("image/")
    if name:
        ext = name.lower().split(".")[-1] if "." in name else ""
        return ext in ["jpg", "jpeg", "png", "gif", "bmp", "webp", "svg"]
    return False


def extract_attachments(value) -> List[Dict]:
    """
    提取附件列表，兼容飞书多维表格附件字段常见结构。
    """
    if not isinstance(value, list):
        return []
    result = []
    for item in value:
        if not isinstance(item, dict):
            continue
        url = item.get("download_url") or item.get("tmp_url") or item.get("url")
        name = item.get("name") or item.get("file_name") or "附件"
        mime = item.get("mime_type") or item.get("type")
        result.append({"name": name, "url": url, "mime": mime})
    return result


def parse_records(raw_records: List[Dict]) -> List[Dict]:
    """
    将飞书接口返回的记录解析为标准结构。
    """
    parsed = []
    for item in raw_records:
        fields = item.get("fields", {})
        subject = fields.get("学科")
        knowledge_points = normalize_to_list(fields.get("知识点"))
        reason_type = normalize_text(fields.get("不会/做错"))
        reason_detail = normalize_text(fields.get("不会/做错原因"))
        
        # 处理"去手写"字段：优先作为附件处理，否则作为文本
        handwriting_raw = fields.get("去手写")
        attachments = extract_attachments(handwriting_raw)
        handwriting_text = ""
        # 如果不是附件列表，则作为文本处理
        if not attachments:
            handwriting_text = normalize_text(handwriting_raw)
        
        # 获取创建时间（飞书API返回的是毫秒时间戳）
        created_time = item.get("created_time", 0)
        if isinstance(created_time, str):
            try:
                created_time = int(created_time)
            except ValueError:
                created_time = 0

        record_id = item.get("record_id") or ""

        parsed.append(
            {
                "record_id": record_id,
                "subject": subject,
                "knowledge_points": knowledge_points,
                "handwriting_text": handwriting_text,
                "reason_type": reason_type,
                "reason_detail": reason_detail,
                "attachments": attachments,
                "created_time": created_time,  # 毫秒时间戳
            }
        )
    return parsed


# ----- 错题练习：练习记录表与选题 -----
# 练习记录表字段名（与飞书多维表格中新建表一致）
_P_FIELD_RID = "错题record_id"
_P_FIELD_LAST = "上次练习时间"
_P_FIELD_MASTERY = "掌握程度"
_P_FIELD_COUNT = "练习次数"
_P_FIELD_NEXT = "下次练习时间"


def _interval_days_for_mastered(n: int) -> int:
    """「会」时按练习次数给出的间隔天数。"""
    return {1: 1, 2: 3, 3: 7, 4: 14}.get(n, 30)


def fetch_practice_records(token: str, practice_table_id: str) -> Dict[str, Dict[str, Any]]:
    """
    拉取练习记录表全部记录，返回 错题record_id -> {practice_record_id, 上次练习时间, 掌握程度, 练习次数, 下次练习时间}。
    同一错题若有多条，保留 上次练习时间 最大的一条。
    """
    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{practice_table_id}/records/search"
    headers = {"Authorization": f"Bearer {token}"}
    page_token = None
    out: Dict[str, Dict[str, Any]] = {}

    while True:
        payload: Dict[str, object] = {"page_size": 100}
        if page_token:
            payload["page_token"] = page_token
        resp = requests.post(url, headers=headers, json=payload, timeout=10)
        if not resp.ok:
            try:
                detail = resp.json()
            except Exception:
                detail = resp.text
            raise RuntimeError(f"拉取练习记录失败 HTTP {resp.status_code}: {detail}")
        data = resp.json()
        if data.get("code") != 0:
            raise RuntimeError(f"拉取练习记录失败: {data}")

        for item in data.get("data", {}).get("items", []):
            fields = item.get("fields", {})
            # 处理 rid 字段，可能是字符串或列表
            rid_raw = fields.get(_P_FIELD_RID)
            if isinstance(rid_raw, list):
                rid = rid_raw[0].strip() if rid_raw and isinstance(rid_raw[0], str) else None
            else:
                rid = (rid_raw or "").strip() or None
            if not rid:
                continue
            try:
                last_ms = int(fields.get(_P_FIELD_LAST) or 0)
            except (TypeError, ValueError):
                last_ms = 0
            try:
                cnt = int(fields.get(_P_FIELD_COUNT) or 0)
            except (TypeError, ValueError):
                cnt = 0
            try:
                next_ms = int(fields.get(_P_FIELD_NEXT) or 0)
            except (TypeError, ValueError):
                next_ms = 0
            # 处理 mastery 字段，可能是字符串或列表
            mastery_raw = fields.get(_P_FIELD_MASTERY)
            if isinstance(mastery_raw, list):
                mastery = mastery_raw[0].strip() if mastery_raw and isinstance(mastery_raw[0], str) else "不会"
            else:
                mastery = (mastery_raw or "").strip() or "不会"

            # 若已存在，只保留 上次练习时间 更大的一条
            if rid in out and (out[rid].get(_P_FIELD_LAST) or 0) >= last_ms:
                continue
            out[rid] = {
                "practice_record_id": item.get("record_id"),
                _P_FIELD_LAST: last_ms,
                _P_FIELD_MASTERY: mastery,
                _P_FIELD_COUNT: cnt,
                _P_FIELD_NEXT: next_ms,
            }

        page_token = data.get("data", {}).get("page_token")
        if not data.get("data", {}).get("has_more"):
            break

    return out


def create_practice_record(
    token: str,
    practice_table_id: str,
    question_record_id: str,
    mastery: str,
    count: int,
    next_ts_ms: int,
) -> Optional[str]:
    """在练习记录表中新建一条记录，返回新记录的 record_id。"""
    now_ms = int(time.time() * 1000)
    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{practice_table_id}/records"
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    body = {
        "fields": {
            _P_FIELD_RID: question_record_id,
            _P_FIELD_LAST: now_ms,
            _P_FIELD_MASTERY: mastery,
            _P_FIELD_COUNT: count,
            _P_FIELD_NEXT: next_ts_ms,
        }
    }
    resp = requests.post(url, headers=headers, json=body, timeout=10)
    if not resp.ok:
        try:
            detail = resp.json()
        except Exception:
            detail = resp.text
        raise RuntimeError(f"创建练习记录失败 HTTP {resp.status_code}: {detail}")
    data = resp.json()
    if data.get("code") != 0:
        raise RuntimeError(f"创建练习记录失败: {data}")
    rec = (data.get("data") or {}).get("record") or {}
    return rec.get("record_id")


def update_practice_record(
    token: str,
    practice_table_id: str,
    practice_record_id: str,
    mastery: str,
    count: int,
    next_ts_ms: int,
) -> None:
    """更新练习记录表中一条记录。"""
    now_ms = int(time.time() * 1000)
    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{APP_TOKEN}/tables/{practice_table_id}/records/{practice_record_id}"
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    body = {
        "fields": {
            _P_FIELD_LAST: now_ms,
            _P_FIELD_MASTERY: mastery,
            _P_FIELD_COUNT: count,
            _P_FIELD_NEXT: next_ts_ms,
        }
    }
    resp = requests.put(url, headers=headers, json=body, timeout=10)
    if not resp.ok:
        try:
            detail = resp.json()
        except Exception:
            detail = resp.text
        raise RuntimeError(f"更新练习记录失败 HTTP {resp.status_code}: {detail}")
    data = resp.json()
    if data.get("code") != 0:
        raise RuntimeError(f"更新练习记录失败: {data}")


def pick_next_question(
    filtered: List[Dict],
    practice_map: Dict[str, Dict[str, Any]],
    now_ms: int,
) -> Optional[Dict]:
    """
    从筛选后的错题中选一道：下次练习时间<=now 优先，否则取下次练习时间最早；无练习记录视为 0 最优先。
    """
    def next_ts(r: Dict) -> int:
        rid = (r.get("record_id") or "").strip()
        if not rid:
            return 0
        p = practice_map.get(rid)
        return int(p.get(_P_FIELD_NEXT, 0) or 0) if p else 0

    # 过滤：至少有 record_id、且 去手写 或 附件 非空
    cand = [r for r in filtered if (r.get("record_id") or "").strip() and (r.get("handwriting_text") or r.get("attachments"))]
    if not cand:
        return None

    # 排序：下次练习时间 <= now 的优先；否则按 下次练习时间 升序
    cand.sort(key=lambda r: (0 if next_ts(r) <= now_ms else 1, next_ts(r)))
    return cand[0]


def save_practice_feedback(
    token: str,
    practice_table_id: str,
    question_record_id: str,
    mastered: bool,
    practice_map: Dict[str, Dict[str, Any]],
) -> None:
    """
    根据用户选择 会/不会 写入或更新练习记录，并就地更新 practice_map 以便本地选题正确。
    mastered=True 表示「会」，False 表示「不会」。
    """
    now_ms = int(time.time() * 1000)
    p = practice_map.get(question_record_id) if question_record_id else None
    prev_count = int(p.get(_P_FIELD_COUNT, 0) or 0) if p else 0
    count = prev_count + 1
    mastery = "会" if mastered else "不会"

    if mastered:
        days = _interval_days_for_mastered(count)
        next_ts_ms = now_ms + days * 24 * 60 * 60 * 1000
    else:
        next_ts_ms = now_ms + 5 * 60 * 1000  # +5 分钟

    if p and p.get("practice_record_id"):
        update_practice_record(token, practice_table_id, p["practice_record_id"], mastery, count, next_ts_ms)
        p[_P_FIELD_LAST] = now_ms
        p[_P_FIELD_MASTERY] = mastery
        p[_P_FIELD_COUNT] = count
        p[_P_FIELD_NEXT] = next_ts_ms
    else:
        new_id = create_practice_record(token, practice_table_id, question_record_id, mastery, count, next_ts_ms)
        practice_map[question_record_id] = {
            "practice_record_id": new_id,
            _P_FIELD_LAST: now_ms,
            _P_FIELD_MASTERY: mastery,
            _P_FIELD_COUNT: count,
            _P_FIELD_NEXT: next_ts_ms,
        }


def _load_image_bytes_for_display(url: str, token: str) -> Optional[bytes]:
    """下载附件图片用于 Streamlit 展示，支持飞书临时 JSON。"""
    if not url or not token:
        return None
    try:
        headers = {"Authorization": f"Bearer {token}"}
        r = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        if not r.ok:
            return None
        ct = (r.headers.get("Content-Type") or "").lower()
        if "application/json" in ct:
            try:
                j = r.json()
                if isinstance(j, dict) and j.get("code") == 0:
                    d = j.get("data", {})
                    u = None
                    tmp = d.get("tmp_download_urls") or []
                    if tmp and isinstance(tmp, list) and tmp:
                        u = tmp[0].get("tmp_download_url") if isinstance(tmp[0], dict) else None
                    if not u:
                        u = d.get("tmp_download_url") or d.get("download_url")
                    if u:
                        r2 = requests.get(u, headers=headers, timeout=15, allow_redirects=True)
                        if r2.ok:
                            return r2.content
            except Exception:
                pass
        return r.content if r.content else None
    except Exception:
        return None


def render_question_streamlit(record: Dict, token: str) -> None:
    """在 Streamlit 中渲染一道题的文本与图片。"""
    t = (record.get("handwriting_text") or "").strip()
    if t:
        st.markdown(t)
    for att in (record.get("attachments") or []):
        if not is_image_file(att.get("name"), att.get("mime")):
            continue
        url = att.get("url")
        raw = _load_image_bytes_for_display(url, token)
        if raw:
            try:
                st.image(io.BytesIO(raw))
            except Exception:
                pass


def safe_get_secret(key: str):
    """
    安全读取 st.secrets，支持两种格式：
    1. 直接格式: KEY = "value"
    2. 嵌套格式: [secrets] 下的 KEY = "value"
    """
    try:
        # 方式1: 直接访问
        value = st.secrets.get(key)
        if value is not None and value != "":
            return value
    except:
        pass
    
    try:
        # 方式2: 从 [secrets] 嵌套结构读取
        if "secrets" in st.secrets:
            value = st.secrets["secrets"].get(key)
            if value is not None and value != "":
                return value
    except:
        pass
    
    return None


def get_config_file_path() -> Path:
    """
    获取配置文件路径（在项目目录下的 .feishu_config.json）。
    """
    return Path(__file__).parent / ".feishu_config.json"


def load_config() -> Dict[str, Optional[str]]:
    """
    从本地配置文件加载凭据。
    """
    config_file = get_config_file_path()
    if config_file.exists():
        try:
            with open(config_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_config(app_id: str, app_secret: str) -> None:
    """
    保存凭据到本地配置文件。
    注意：在 Streamlit Cloud 等只读文件系统上，此操作会静默失败。
    """
    config_file = get_config_file_path()
    try:
        config = {"FEISHU_APP_ID": app_id, "FEISHU_APP_SECRET": app_secret}
        with open(config_file, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2)
        # 设置文件权限（仅所有者可读写）
        if os.name != "nt":  # 非Windows系统
            os.chmod(config_file, 0o600)
    except (IOError, OSError, PermissionError):
        # 在 Streamlit Cloud 等只读文件系统上，保存失败是正常的
        # 配置应通过环境变量或 st.secrets 提供
        pass
    except Exception:
        pass  # 其他错误也静默处理


def build_doc(subjects: List[str], selections: Dict[str, List[Dict]], token: str) -> bytes:
    """
    根据选择生成 Word 文档二进制内容。
    """
    doc = Document()
    title = "、".join(subjects) if subjects else "错题"
    doc.add_heading(f"{title} 错题专项训练", 0)

    for kp, questions in selections.items():
        # 跳过空列表（生成失败的题目）
        if not questions:
            continue
        doc.add_heading(kp, level=1)
        for idx, q in enumerate(questions, start=1):
            # 题目内容：优先使用附件（图片），否则使用文本
            attachments = q.get("attachments") or []
            handwriting_text = q.get("handwriting_text", "").strip()
            
            if attachments:
                # 有附件，使用 List Number 编号；第一段内容放入 para，避免空编号
                para = doc.add_paragraph(style="List Number")
                first = True

                for att in attachments:
                    url = att.get("url")
                    name = att.get("name") or "附件"
                    mime = att.get("mime")
                    if not url:
                        continue

                    is_image = is_image_file(name, mime)

                    if not is_image:
                        text = f"附件：{name}（非图片，下载链接：{url}）"
                        if first:
                            r = para.add_run(text)
                            r.italic = True
                            first = False
                        else:
                            p = doc.add_paragraph(text)
                            if p.runs:
                                p.runs[0].italic = True
                        continue

                    try:
                        headers = {"Authorization": f"Bearer {token}"}
                        resp = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
                        if not resp.ok:
                            text = f"[附件下载失败] {name} - HTTP {resp.status_code}"
                            if first:
                                para.add_run(text)
                                first = False
                            else:
                                doc.add_paragraph(text)
                            continue

                        content_type = resp.headers.get("Content-Type", "").lower()
                        image_data = None

                        if "application/json" in content_type:
                            try:
                                json_data = resp.json()
                                if isinstance(json_data, dict) and json_data.get("code") == 0:
                                    data = json_data.get("data", {})
                                    tmp_urls = data.get("tmp_download_urls", [])
                                    if tmp_urls and isinstance(tmp_urls, list) and len(tmp_urls) > 0:
                                        real_url = tmp_urls[0].get("tmp_download_url") if isinstance(tmp_urls[0], dict) else None
                                    else:
                                        real_url = data.get("tmp_download_url") or data.get("download_url") or json_data.get("download_url")

                                    if real_url:
                                        resp2 = requests.get(real_url, headers=headers, timeout=15, allow_redirects=True)
                                        if resp2.ok:
                                            image_data = resp2.content
                                        else:
                                            text = f"[附件下载失败] {name} - HTTP {resp2.status_code}"
                                            if first:
                                                para.add_run(text)
                                                first = False
                                            else:
                                                doc.add_paragraph(text)
                                            continue
                                    else:
                                        text = f"[无法获取附件下载地址] {name}"
                                        if first:
                                            para.add_run(text)
                                            first = False
                                        else:
                                            doc.add_paragraph(text)
                                        continue
                            except (ValueError, KeyError, TypeError):
                                image_data = resp.content
                        else:
                            image_data = resp.content

                        if image_data:
                            try:
                                image_stream = io.BytesIO(image_data)
                                if first:
                                    run = para.add_run()
                                    run.add_picture(image_stream, width=Inches(5.5))
                                    first = False
                                else:
                                    doc.add_picture(image_stream, width=Inches(5.5))
                            except Exception as img_exc:  # noqa: BLE001
                                text = f"附件：{name}（图片插入失败：{img_exc}）"
                                if first:
                                    r = para.add_run(text)
                                    r.italic = True
                                    first = False
                                else:
                                    p = doc.add_paragraph(text)
                                    if p.runs:
                                        p.runs[0].italic = True
                        else:
                            text = f"[附件处理失败] {name}"
                            if first:
                                para.add_run(text)
                                first = False
                            else:
                                doc.add_paragraph(text)
                    except Exception as exc:  # noqa: BLE001
                        text = f"[附件处理异常] {name}: {exc}"
                        if first:
                            para.add_run(text)
                            first = False
                        else:
                            doc.add_paragraph(text)
            elif handwriting_text:
                # 没有附件但有文本，显示文本
                para = doc.add_paragraph(style="List Number")
                para.add_run(handwriting_text)
            else:
                # 既无附件也无文本
                para = doc.add_paragraph(style="List Number")
                para.add_run("（无题干）")
            
            # 追加备注信息，便于回顾错因
            if q.get("reason_type") or q.get("reason_detail"):
                note = f"错因：{q.get('reason_type') or ''} {q.get('reason_detail') or ''}".strip()
                if note.startswith("错因："):
                    doc.add_paragraph(note).italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_html(subjects: List[str], selections: Dict[str, List[Dict]], token: str) -> str:
    """
    根据选择生成 HTML 文档内容。
    """
    title = "、".join(subjects) if subjects else "错题"
    
    html_parts = [
        "<!DOCTYPE html>",
        "<html lang='zh-CN'>",
        "<head>",
        "    <meta charset='UTF-8'>",
        "    <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
        f"    <title>{title} 错题专项训练</title>",
        "    <style>",
        "        body { font-family: 'Microsoft YaHei', Arial, sans-serif; line-height: 1.6; padding: 20px; max-width: 1200px; margin: 0 auto; }",
        "        h1 { color: #333; border-bottom: 3px solid #4CAF50; padding-bottom: 10px; }",
        "        h2 { color: #555; margin-top: 30px; border-left: 5px solid #2196F3; padding-left: 15px; }",
        "        .question { margin: 20px 0; padding: 15px; background: #f9f9f9; border-radius: 5px; }",
        "        .question-number { font-weight: bold; color: #2196F3; margin-right: 10px; }",
        "        .question-content { margin: 10px 0; }",
        "        .question-content img { max-width: 100%; height: auto; margin: 10px 0; border: 1px solid #ddd; border-radius: 5px; }",
        "        .reason { font-style: italic; color: #666; margin-top: 10px; padding-left: 20px; }",
        "        .error-note { color: #ff9800; font-size: 0.9em; }",
        "    </style>",
        "</head>",
        "<body>",
        f"    <h1>{title} 错题专项训练</h1>",
    ]
    
    for kp, questions in selections.items():
        # 跳过空列表（生成失败的题目）
        if not questions:
            continue
        html_parts.append(f"    <h2>{kp}</h2>")
        
        for idx, q in enumerate(questions, start=1):
            html_parts.append("    <div class='question'>")
            html_parts.append(f"        <div class='question-number'>{idx}.</div>")
            html_parts.append("        <div class='question-content'>")
            
            attachments = q.get("attachments") or []
            handwriting_text = q.get("handwriting_text", "").strip()
            
            if attachments:
                # 处理附件（图片）
                for att in attachments:
                    url = att.get("url")
                    name = att.get("name") or "附件"
                    mime = att.get("mime")
                    
                    if not url:
                        continue
                    
                    is_image = is_image_file(name, mime)
                    
                    if is_image:
                        try:
                            headers = {"Authorization": f"Bearer {token}"}
                            resp = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
                            
                            if resp.ok:
                                content_type = resp.headers.get("Content-Type", "").lower()
                                image_data = None
                                final_content_type = content_type
                                
                                # 处理JSON响应
                                if "application/json" in content_type:
                                    try:
                                        json_data = resp.json()
                                        if isinstance(json_data, dict) and json_data.get("code") == 0:
                                            data = json_data.get("data", {})
                                            tmp_urls = data.get("tmp_download_urls", [])
                                            if tmp_urls and isinstance(tmp_urls, list) and len(tmp_urls) > 0:
                                                real_url = tmp_urls[0].get("tmp_download_url") if isinstance(tmp_urls[0], dict) else None
                                            else:
                                                real_url = data.get("tmp_download_url") or data.get("download_url")
                                            
                                            if real_url:
                                                resp2 = requests.get(real_url, headers=headers, timeout=15, allow_redirects=True)
                                                if resp2.ok:
                                                    image_data = resp2.content
                                                    final_content_type = resp2.headers.get("Content-Type", "image/png").lower()
                                    except Exception:
                                        pass
                                
                                if not image_data:
                                    image_data = resp.content
                                    final_content_type = content_type
                                
                                # 转换为base64嵌入
                                img_base64 = base64.b64encode(image_data).decode('utf-8')
                                img_src = f"data:{final_content_type or 'image/png'};base64,{img_base64}"
                                html_parts.append(f"            <img src='{img_src}' alt='{name}' />")
                            else:
                                html_parts.append(f"            <p class='error-note'>[图片加载失败] {name}</p>")
                        except Exception as exc:
                            html_parts.append(f"            <p class='error-note'>[图片加载异常] {name}: {exc}</p>")
                    else:
                        html_parts.append(f"            <p>附件：<a href='{url}' target='_blank'>{name}</a></p>")
            elif handwriting_text:
                # 显示文本内容（转义HTML特殊字符）
                escaped_text = html.escape(handwriting_text)
                html_parts.append(f"            <div>{escaped_text.replace(chr(10), '<br>')}</div>")
            else:
                html_parts.append("            <div>（无题干）</div>")
            
            html_parts.append("        </div>")
            
            # 添加错因备注
            if q.get("reason_type") or q.get("reason_detail"):
                reason = f"{q.get('reason_type') or ''} {q.get('reason_detail') or ''}".strip()
                if reason:
                    html_parts.append(f"        <div class='reason'>错因：{html.escape(reason)}</div>")
            
            html_parts.append("    </div>")
    
    html_parts.extend([
        "</body>",
        "</html>"
    ])
    
    return "\n".join(html_parts)


def generate_similar_questions_with_llm(reference_question: Dict, count: int, api_key: str, api_base: str = None, model: str = None, token: str = None) -> List[str]:
    """
    使用大模型生成类似题目。
    
    Args:
        reference_question: 参考题目（包含handwriting_text或attachments）
        count: 需要生成的题目数量
        api_key: API密钥
        api_base: API基础URL（智谱AI API Base URL）
        model: 模型名称（可选，如果不指定则根据api_base自动选择）
    
    Returns:
        生成的题目列表
    """
    # 构建参考题目的文本描述
    ref_text = reference_question.get("handwriting_text", "").strip()
    attachments = reference_question.get("attachments", [])
    
    if not ref_text and not attachments:
        raise ValueError("参考题目不能为空")
    
    # 检查是否有图片附件
    image_attachments = [att for att in attachments if is_image_file(att.get("name", ""), att.get("mime"))]
    has_images = len(image_attachments) > 0
    
    # 构建提示词
    if ref_text:
        question_description = ref_text
    elif has_images:
        question_description = "（请查看图片中的题目内容）"
    else:
        question_description = "[题目内容]"
    
    prompt_text = f"""你是一位经验丰富的教师，需要基于参考题目生成类似的新题目。

参考题目：{question_description}

请生成 {count} 道类似的题目，要求：

1. **保持核心要素一致**：
   - 保持相同的知识点和解题方法
   - 保持相同的题目类型（如选择题、填空题、计算题等）
   - 保持相同的难度级别
   - 如果是数学题，保持相同的运算类型和公式结构

2. **只改变可变要素**：
   - 可以改变具体数字、数值（如：把"5+3"改为"7+4"）
   - 可以改变具体的人物、物品、场景名称
   - 可以改变题目的表述方式，但核心意思保持一致
   - 保持题目的结构和解题步骤一致

3. **输出格式**：
   - 每道题目单独一行
   - 只输出题目内容，不要编号，不要添加"题目1"、"题目2"等前缀
   - 不要添加任何解释说明
   - 确保每道题目都是完整、独立、可以直接使用的

4. **质量要求**：
   - 题目必须合理、可解，不能出现逻辑错误
   - 题目必须与原题难度相当
   - 不能生成完全相同的题目，但也不能偏离太远
   - 生成的题目应该可以直接用于练习

请严格按照以上要求生成 {count} 道类似题目，每行一道："""
    
    # 默认使用智谱AI GLM-4.6V（支持多模态）
    if not model:
        model = "glm-4.6v"
    
    # 调用智谱AI API
    # 智谱AI的API URL格式
    if api_base and api_base.endswith("/chat/completions"):
        api_url = api_base
    else:
        api_url = (api_base or "https://open.bigmodel.cn/api/paas/v4").rstrip('/') + "/chat/completions"
    # 智谱AI的认证格式
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # 检查模型是否支持多模态（图片输入）
    model_lower = (model or "").lower()
    supports_vision = (
        "4.6v" in model_lower or 
        "glm-4-6v" in model_lower or 
        "glm-4.6v" in model_lower or
        "vision" in model_lower or 
        "4o" in model_lower or
        "gpt-4o" in model_lower
    )
    
    # 构建消息内容
    if has_images and supports_vision:
        # 构建多模态消息（包含图片）
        content_list = []
        image_added = False
        
        # 添加图片（使用缓存）
        for img_att in image_attachments[:1]:  # 只使用第一张图片
            img_url = img_att.get("url")
            if img_url and token:
                # 使用缓存获取图片
                cached = _get_cached_image_base64(img_url, token)
                if cached:
                    img_base64, img_mime = cached
                    content_list.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{img_mime};base64,{img_base64}"
                        }
                    })
                    image_added = True
        
        # 添加文本提示
        content_list.append({
            "type": "text",
            "text": prompt_text
        })
        
        # 如果成功添加了图片，使用多模态消息；否则回退到纯文本
        if image_added:
            messages = [{"role": "user", "content": content_list}]
        else:
            # 图片添加失败，如果有文本内容，使用文本；否则抛出错误
            if ref_text:
                # 有文本内容，重新构建提示词（不提及图片）
                fallback_prompt = f"""你是一位经验丰富的教师，需要基于参考题目生成类似的新题目。

参考题目：{ref_text}

请生成 {count} 道类似的题目，要求：

1. **保持核心要素一致**：
   - 保持相同的知识点和解题方法
   - 保持相同的题目类型（如选择题、填空题、计算题等）
   - 保持相同的难度级别
   - 如果是数学题，保持相同的运算类型和公式结构

2. **只改变可变要素**：
   - 可以改变具体数字、数值（如：把"5+3"改为"7+4"）
   - 可以改变具体的人物、物品、场景名称
   - 可以改变题目的表述方式，但核心意思保持一致
   - 保持题目的结构和解题步骤一致

3. **输出格式**：
   - 每道题目单独一行
   - 只输出题目内容，不要编号，不要添加"题目1"、"题目2"等前缀
   - 不要添加任何解释说明
   - 确保每道题目都是完整、独立、可以直接使用的

4. **质量要求**：
   - 题目必须合理、可解，不能出现逻辑错误
   - 题目必须与原题难度相当
   - 不能生成完全相同的题目，但也不能偏离太远
   - 生成的题目应该可以直接用于练习

请严格按照以上要求生成 {count} 道类似题目，每行一道："""
                messages = [{"role": "user", "content": fallback_prompt}]
            else:
                # 没有文本也没有图片，无法生成题目
                raise ValueError("题目包含图片但图片处理失败，且没有文本内容，无法生成类似题目。请检查图片URL或网络连接。")
    else:
        # 纯文本消息
        messages = [{"role": "user", "content": prompt_text}]
    
    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(api_url, headers=headers, json=payload, timeout=60)
        
        # 检查响应状态
        if not response.ok:
            # 获取详细的错误信息
            try:
                error_detail = response.json()
                error_msg = f"API错误 {response.status_code}: {error_detail}"
            except:
                error_msg = f"API错误 {response.status_code}: {response.text[:200]}"
            
            # 根据不同错误码提供更详细的提示
            if response.status_code == 400:
                error_msg += f"\n请求URL: {api_url}\n模型: {model}\n请检查模型名称、API Key和请求格式是否正确。"
            elif response.status_code == 401:
                error_msg += f"\n请求URL: {api_url}\n模型: {model}\n⚠️ API Key无效或已过期，请检查：\n1. API Key是否正确\n2. API Key是否已过期\n3. API Key是否有足够的权限访问该模型"
            
            raise Exception(error_msg)
        
        response.raise_for_status()
        result = response.json()
        
        # 检查响应格式
        if "choices" not in result or not result["choices"]:
            raise Exception(f"API响应格式错误: {result}")
        
        # 提取生成的文本
        generated_text = result["choices"][0]["message"]["content"].strip()
        
        # 按行分割，过滤空行
        questions = [q.strip() for q in generated_text.split("\n") if q.strip()]
        
        # 如果生成的数量不够，重复最后一道题
        while len(questions) < count:
            questions.append(questions[-1] if questions else "（生成失败）")
        
        # 如果生成的数量太多，只取前count个
        return questions[:count]
    
    except Exception as e:
        # 如果API调用失败，抛出异常以便上层处理
        raise Exception(f"题目生成失败: {str(e)}")


def _load_app_config():
    """加载应用配置，返回 (app_id, app_secret, llm_api_key, llm_api_base, llm_model, config, is_streamlit_cloud)"""
    config = load_config()
    
    # 检测是否在 Streamlit Cloud 上运行
    try:
        _ = st.secrets
        is_streamlit_cloud = True
    except StreamlitSecretNotFoundError:
        is_streamlit_cloud = False
    except (AttributeError, RuntimeError, Exception):
        is_streamlit_cloud = False
    
    # 读取飞书配置
    env_app_id = os.getenv("FEISHU_APP_ID")
    env_app_secret = os.getenv("FEISHU_APP_SECRET")
    secret_app_id = safe_get_secret("FEISHU_APP_ID")
    secret_app_secret = safe_get_secret("FEISHU_APP_SECRET")
    config_app_id = config.get("FEISHU_APP_ID")
    config_app_secret = config.get("FEISHU_APP_SECRET")
    session_app_id = st.session_state.get("feishu_app_id")
    session_app_secret = st.session_state.get("feishu_app_secret")
    
    app_id = env_app_id or secret_app_id or config_app_id or session_app_id
    app_secret = env_app_secret or secret_app_secret or config_app_secret or session_app_secret
    
    # 读取LLM配置
    llm_api_key = (
        os.getenv("LLM_API_KEY")
        or safe_get_secret("LLM_API_KEY")
        or config.get("LLM_API_KEY")
        or st.session_state.get("llm_api_key")
    )
    llm_api_base = "https://open.bigmodel.cn/api/paas/v4"  # 固定使用智谱AI
    llm_model = (
        os.getenv("LLM_MODEL")
        or safe_get_secret("LLM_MODEL")
        or config.get("LLM_MODEL")
        or st.session_state.get("llm_model")
        or "glm-4.6v"
    )
    
    return app_id, app_secret, llm_api_key, llm_api_base, llm_model, config, is_streamlit_cloud


def _check_feishu_credentials(app_id, app_secret, is_streamlit_cloud):
    """检查飞书凭据，如果缺失则显示配置界面"""
    if not app_id or not app_secret:
        if is_streamlit_cloud:
            missing_items = []
            if not app_id:
                missing_items.append("FEISHU_APP_ID")
            if not app_secret:
                missing_items.append("FEISHU_APP_SECRET")
            st.error(f"❌ 配置缺失：{', '.join(missing_items)}，请在 Streamlit Cloud Secrets 中配置。")
            st.stop()
        else:
            st.info("请输入飞书应用凭据")
            app_id_input = st.text_input("FEISHU_APP_ID", value=app_id or "")
            app_secret_input = st.text_input("FEISHU_APP_SECRET", value=app_secret or "", type="password")
            if not app_id_input or not app_secret_input:
                st.stop()
            save_config(app_id_input, app_secret_input)
            st.session_state["feishu_app_id"] = app_id_input
            st.session_state["feishu_app_secret"] = app_secret_input
            return app_id_input, app_secret_input
    return app_id, app_secret


def _render_home_page():
    """渲染主页：两个大按钮"""
    st.title("📚 错题本")
    st.caption(f"v{VERSION}")
    
    st.markdown("---")
    
    # 使用容器创建更美观的按钮布局
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div style="text-align: center; padding: 20px;">
            <h2>📝</h2>
            <p>根据艾宾浩斯遗忘曲线复习错题</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("错题练习", type="primary", use_container_width=True, key="home_practice_btn"):
            st.session_state["current_page"] = "practice"
            st.rerun()
    
    with col2:
        st.markdown("""
        <div style="text-align: center; padding: 20px;">
            <h2>📄</h2>
            <p>选择学科和知识点生成试卷</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("生成试卷", type="primary", use_container_width=True, key="home_exam_btn"):
            st.session_state["current_page"] = "exam"
            st.rerun()


# ==================== 练习优化相关函数 ====================

def _get_today_str() -> str:
    """获取今天的日期字符串"""
    return datetime.now().strftime("%Y-%m-%d")


def _init_daily_practice_tracking():
    """初始化每日练习追踪，每天重置"""
    today = _get_today_str()
    if st.session_state.get("practice_date") != today:
        st.session_state["practiced_today"] = set()
        st.session_state["practice_date"] = today
        st.session_state["similar_cache"] = {}  # 每天也清空缓存
        st.session_state["pregenerate_queue"] = []
        st.session_state["pregenerate_done"] = set()


def _mark_practiced_today(record_id: str):
    """标记某题今日已练过"""
    if not record_id:
        return
    _init_daily_practice_tracking()
    practiced = st.session_state.get("practiced_today", set())
    practiced.add(record_id)
    st.session_state["practiced_today"] = practiced


def _is_practiced_today(record_id: str) -> bool:
    """检查某题今日是否已练过"""
    if not record_id:
        return False
    _init_daily_practice_tracking()
    return record_id in st.session_state.get("practiced_today", set())


def _filter_not_practiced_today(questions: List[Dict]) -> List[Dict]:
    """过滤掉今日已练过的题目"""
    _init_daily_practice_tracking()
    practiced = st.session_state.get("practiced_today", set())
    return [q for q in questions if (q.get("record_id") or "").strip() not in practiced]


def _get_similar_from_cache(record_id: str) -> Optional[str]:
    """从缓存获取类似题"""
    if not record_id:
        return None
    cache = st.session_state.get("similar_cache", {})
    if record_id in cache and cache[record_id]:
        # 取出一道（不删除，因为可能需要第二道）
        return cache[record_id][0] if cache[record_id] else None
    return None


def _get_second_similar_from_cache(record_id: str) -> Optional[str]:
    """从缓存获取第二道类似题"""
    if not record_id:
        return None
    cache = st.session_state.get("similar_cache", {})
    if record_id in cache and len(cache[record_id]) >= 2:
        return cache[record_id][1]
    return None


def _add_to_similar_cache(record_id: str, similar_texts: List[str]):
    """添加类似题到缓存"""
    if not record_id or not similar_texts:
        return
    if "similar_cache" not in st.session_state:
        st.session_state["similar_cache"] = {}
    st.session_state["similar_cache"][record_id] = similar_texts


def _get_cached_image_base64(img_url: str, token: str) -> Optional[tuple]:
    """
    获取图片的base64编码，优先从缓存读取
    返回 (base64_data, mime_type) 或 None
    """
    if not img_url:
        return None
    
    # 检查缓存
    if "image_cache" not in st.session_state:
        st.session_state["image_cache"] = {}
    
    cache = st.session_state["image_cache"]
    if img_url in cache:
        return cache[img_url]
    
    # 缓存未命中，下载图片
    try:
        img_headers = {"Authorization": f"Bearer {token}"}
        img_resp = requests.get(img_url, headers=img_headers, timeout=15, allow_redirects=True)
        
        if img_resp.ok:
            content_type = img_resp.headers.get("Content-Type", "").lower()
            image_data = None
            
            if "application/json" in content_type:
                try:
                    json_data = img_resp.json()
                    if isinstance(json_data, dict) and json_data.get("code") == 0:
                        data = json_data.get("data", {})
                        tmp_urls = data.get("tmp_download_urls", [])
                        if tmp_urls and isinstance(tmp_urls, list) and len(tmp_urls) > 0:
                            real_url = tmp_urls[0].get("tmp_download_url") if isinstance(tmp_urls[0], dict) else None
                        else:
                            real_url = data.get("tmp_download_url") or data.get("download_url")
                        
                        if real_url:
                            img_resp2 = requests.get(real_url, headers=img_headers, timeout=15, allow_redirects=True)
                            if img_resp2.ok:
                                image_data = img_resp2.content
                                content_type = img_resp2.headers.get("Content-Type", "image/png").lower()
                except Exception:
                    pass
            
            if not image_data:
                image_data = img_resp.content
            
            if image_data and len(image_data) > 0:
                img_base64 = base64.b64encode(image_data).decode('utf-8')
                img_mime = content_type if content_type and "image" in content_type else "image/png"
                result = (img_base64, img_mime)
                cache[img_url] = result
                return result
    except Exception:
        pass
    
    return None


def _pregenerate_one_similar(question: Dict, llm_api_key: str, llm_api_base: str, llm_model: str, token: str) -> bool:
    """为一道题预生成类似题，返回是否成功"""
    record_id = (question.get("record_id") or "").strip()
    if not record_id:
        return False
    
    # 已经生成过则跳过
    done = st.session_state.get("pregenerate_done", set())
    if record_id in done:
        return True
    
    try:
        # 生成2道类似题（第一次不会和第二次不会各用一道）
        texts = generate_similar_questions_with_llm(question, 2, llm_api_key, llm_api_base, llm_model, token)
        if texts:
            _add_to_similar_cache(record_id, texts)
            done.add(record_id)
            st.session_state["pregenerate_done"] = done
            return True
    except Exception:
        pass
    return False


def _get_pregenerate_progress() -> tuple:
    """获取预生成进度 (已完成, 总数)"""
    done = len(st.session_state.get("pregenerate_done", set()))
    queue = st.session_state.get("pregenerate_queue", [])
    total = len(queue)
    return (done, total)


def _render_practice_page(token, records, llm_api_key, llm_api_base, llm_model, config):
    """渲染错题练习页面"""
    # 初始化每日练习追踪
    _init_daily_practice_tracking()
    
    # 返回按钮
    if st.button("← 返回主页", key="practice_back"):
        # 清理练习状态
        for k in ("practice_current", "practice_origin", "practice_is_similar", "practice_similar_count", "practice_map", "practice_filtered", "practice_table_id", "pregenerate_queue", "pregenerate_done"):
            st.session_state.pop(k, None)
        st.session_state["current_page"] = "home"
        st.rerun()
    
    st.title("📝 错题练习")
    st.caption("根据艾宾浩斯遗忘曲线智能安排复习（错题原题每天只出现一次）")
    
    practice_table_id = (
        os.getenv("FEISHU_PRACTICE_TABLE_ID")
        or safe_get_secret("FEISHU_PRACTICE_TABLE_ID")
        or config.get("FEISHU_PRACTICE_TABLE_ID")
        or ""
    )
    
    if not practice_table_id:
        st.error(
            "错题练习需要配置 **FEISHU_PRACTICE_TABLE_ID**（练习记录表的 table_id）。\n\n"
            "请在 `.feishu_config.json` 中配置。\n\n"
            "需在同一多维表格下新建一张表，包含字段：错题record_id、上次练习时间、掌握程度、练习次数、下次练习时间。"
        )
        return
    
    # 学科筛选
    subjects = sorted({r["subject"] for r in records if r.get("subject")})
    selected_subjects = st.multiselect("选择学科", options=subjects, default=subjects, key="practice_subjects")
    filtered = [r for r in records if r.get("subject") in selected_subjects]
    
    # 知识点筛选
    knowledge_options = sorted({kp for r in filtered for kp in r.get("knowledge_points") or []})
    selected_kp = st.multiselect("选择知识点", options=knowledge_options, default=knowledge_options, key="practice_kp")
    
    filtered_practice = [r for r in filtered if any(kp in (r.get("knowledge_points") or []) for kp in selected_kp)] if selected_kp else filtered
    
    st.markdown("---")
    
    # 显示预生成进度
    done_count, total_count = _get_pregenerate_progress()
    if total_count > 0:
        if done_count < total_count:
            st.caption(f"⏳ 正在准备类似题... ({done_count}/{total_count})")
        else:
            st.caption(f"✓ 类似题已就绪 ({done_count}/{total_count})")
    
    def _go_next_practice() -> None:
        """进入下一道题，并标记当前题已练过"""
        # 标记当前题今日已练
        cur = st.session_state.get("practice_current")
        if cur and not st.session_state.get("practice_is_similar"):
            rid = (cur.get("record_id") or "").strip()
            if rid:
                _mark_practiced_today(rid)
        
        pm = st.session_state.get("practice_map", {})
        pf = st.session_state.get("practice_filtered", [])
        
        # 过滤掉今日已练过的题目
        pf_available = _filter_not_practiced_today(pf)
        
        n = pick_next_question(pf_available, pm, int(time.time() * 1000))
        if n:
            st.session_state["practice_current"] = n
            st.session_state["practice_origin"] = None
            st.session_state["practice_is_similar"] = False
            st.session_state["practice_similar_count"] = 0
        else:
            for k in ("practice_current", "practice_origin", "practice_is_similar", "practice_similar_count"):
                st.session_state.pop(k, None)
            st.success("🎉 本轮可复习的题目已练完！")
    
    if st.session_state.get("practice_current"):
        cur = st.session_state["practice_current"]
        st.session_state.setdefault("practice_map", {})
        st.session_state.setdefault("practice_filtered", [])
        
        # 显示题目
        st.markdown("### 当前题目")
        if st.session_state.get("practice_is_similar"):
            st.caption("📌 类似题")
        
        render_question_streamlit(cur, token)
        
        st.markdown("---")
        st.markdown("**掌握了吗？**")
        
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("✓ 会了", type="primary", use_container_width=True, key="practice_btn_yes"):
                is_sim = st.session_state.get("practice_is_similar", False)
                if not is_sim:
                    save_practice_feedback(
                        token,
                        st.session_state["practice_table_id"],
                        (cur.get("record_id") or "").strip(),
                        True,
                        st.session_state["practice_map"],
                    )
                _go_next_practice()
                st.rerun()
        
        with col_b:
            if st.button("✗ 不会", use_container_width=True, key="practice_btn_no"):
                is_sim = st.session_state.get("practice_is_similar", False)
                orig = st.session_state.get("practice_origin")
                ptid = st.session_state.get("practice_table_id", "")
                pm = st.session_state.get("practice_map", {})
                
                if not is_sim:
                    # 第一次点击"不会"
                    rid = (cur.get("record_id") or "").strip()
                    save_practice_feedback(token, ptid, rid, False, pm)
                    st.session_state["practice_origin"] = cur
                    
                    # 优先从缓存获取类似题
                    cached_similar = _get_similar_from_cache(rid)
                    if cached_similar:
                        st.session_state["practice_current"] = {"handwriting_text": cached_similar, "attachments": [], "record_id": ""}
                        st.session_state["practice_is_similar"] = True
                        st.session_state["practice_similar_count"] = 1
                        st.rerun()
                    elif llm_api_key:
                        # 缓存未命中，实时生成
                        with st.spinner("正在生成类似题目…"):
                            try:
                                texts = generate_similar_questions_with_llm(cur, 2, llm_api_key, llm_api_base, llm_model, token)
                                if texts:
                                    _add_to_similar_cache(rid, texts)
                                    st.session_state["practice_current"] = {"handwriting_text": texts[0], "attachments": [], "record_id": ""}
                                    st.session_state["practice_is_similar"] = True
                                    st.session_state["practice_similar_count"] = 1
                                    st.rerun()
                                else:
                                    _go_next_practice()
                            except Exception as e:
                                st.error(f"生成类似题目失败：{e}")
                                _go_next_practice()
                    else:
                        _go_next_practice()
                else:
                    # 第二次点击"不会"（在类似题上）
                    cnt = st.session_state.get("practice_similar_count", 0)
                    if cnt < 2 and orig:
                        orig_rid = (orig.get("record_id") or "").strip()
                        # 优先从缓存获取第二道类似题
                        cached_second = _get_second_similar_from_cache(orig_rid)
                        if cached_second:
                            st.session_state["practice_current"] = {"handwriting_text": cached_second, "attachments": [], "record_id": ""}
                            st.session_state["practice_similar_count"] = 2
                            st.rerun()
                        elif llm_api_key:
                            with st.spinner("再出一道类似题目…"):
                                try:
                                    texts = generate_similar_questions_with_llm(orig, 1, llm_api_key, llm_api_base, llm_model, token)
                                    if texts:
                                        st.session_state["practice_current"] = {"handwriting_text": texts[0], "attachments": [], "record_id": ""}
                                        st.session_state["practice_similar_count"] = 2
                                        st.rerun()
                                    else:
                                        _go_next_practice()
                            except Exception:
                                _go_next_practice()
                    else:
                        _go_next_practice()
                st.rerun()
    else:
        st.info("点击下方按钮开始练习")
        if st.button("🚀 开始练习", type="primary", use_container_width=True, key="practice_start"):
            with st.spinner("正在加载练习记录…"):
                try:
                    pm = fetch_practice_records(token, practice_table_id)
                    
                    # 过滤掉今日已练过的题目
                    available_questions = _filter_not_practiced_today(filtered_practice)
                    
                    n = pick_next_question(available_questions, pm, int(time.time() * 1000))
                    if not n:
                        st.info("暂无需要复习的题目，或今日的题目已全部练完。")
                    else:
                        # 立即显示第一道题
                        st.session_state["practice_current"] = n
                        st.session_state["practice_map"] = pm
                        st.session_state["practice_table_id"] = practice_table_id
                        st.session_state["practice_filtered"] = filtered_practice
                        st.session_state["practice_origin"] = None
                        st.session_state["practice_is_similar"] = False
                        st.session_state["practice_similar_count"] = 0
                        
                        # 设置预生成队列（所有可练习的题目）
                        st.session_state["pregenerate_queue"] = available_questions
                        st.session_state["pregenerate_done"] = set()
                        st.session_state["pregenerate_started"] = True
                        
                        st.rerun()
                except Exception as e:
                    st.error(f"加载练习记录失败：{e}")
    
    # 后台预生成逻辑：每次页面刷新时尝试生成一道
    if st.session_state.get("pregenerate_started") and llm_api_key:
        queue = st.session_state.get("pregenerate_queue", [])
        done = st.session_state.get("pregenerate_done", set())
        
        # 找到下一个需要预生成的题目
        for q in queue:
            rid = (q.get("record_id") or "").strip()
            if rid and rid not in done:
                # 预生成这道题的类似题（不阻塞UI）
                _pregenerate_one_similar(q, llm_api_key, llm_api_base, llm_model, token)
                break  # 每次只生成一道，避免阻塞太久
    
    # 底部返回按钮
    st.markdown("---")
    if st.button("← 返回主页", key="practice_back_bottom"):
        for k in ("practice_current", "practice_origin", "practice_is_similar", "practice_similar_count", "practice_map", "practice_filtered", "practice_table_id", "pregenerate_queue", "pregenerate_done", "pregenerate_started"):
            st.session_state.pop(k, None)
        st.session_state["current_page"] = "home"
        st.rerun()


def _render_exam_page(token, records, llm_api_key, llm_api_base, llm_model):
    """渲染生成试卷页面"""
    # 返回按钮
    if st.button("← 返回主页", key="exam_back"):
        st.session_state["current_page"] = "home"
        st.rerun()
    
    st.title("📄 生成试卷")
    st.caption("选择学科和知识点，生成错题专项训练")
    
    # 学科选择
    subjects = sorted({r["subject"] for r in records if r.get("subject")})
    if not subjects:
        st.warning("没有找到学科数据")
        return
    
    selected_subjects = st.multiselect("选择学科", options=subjects, default=subjects, key="exam_subjects")
    if not selected_subjects:
        st.info("请选择至少一个学科")
        return
    
    filtered = [r for r in records if r.get("subject") in selected_subjects]
    
    # 知识点选择
    knowledge_options = sorted({kp for r in filtered for kp in r.get("knowledge_points") or []})
    selected_kp = st.multiselect("选择知识点", options=knowledge_options, default=knowledge_options, key="exam_kp")
    
    # 每个知识点的题目数量
    selected_plan: Dict[str, int] = {}
    for kp in selected_kp:
        pool = [r for r in filtered if kp in (r.get("knowledge_points") or [])]
        max_count = len(pool)
        count = st.number_input(f"{kp}（最多 {max_count} 题）", min_value=0, max_value=max_count, value=max_count, key=f"exam_count_{kp}")
        selected_plan[kp] = count
    
    # 显示当前选择的总题目数量
    total_count = sum(selected_plan.values())
    st.markdown(f"### 当前选择题目数量：{total_count} 道")
    
    has_valid_selection = any(count > 0 for count in selected_plan.values())
    
    if not has_valid_selection:
        st.info("请至少选择一道题目")
        return
    
    def prepare_selections():
        selections: Dict[str, List[Dict]] = {}
        for kp, count in selected_plan.items():
            if count <= 0:
                continue
            pool = [r for r in filtered if kp in (r.get("knowledge_points") or [])]
            if count > len(pool):
                count = len(pool)
            if count > 0:
                selections[kp] = random.sample(pool, count)
        return selections
    
    def prepare_similar_selections():
        similar_selections: Dict[str, List[Dict]] = {}
        total_upper = sum(c for c in selected_plan.values() if c > 0)
        progress_bar = st.progress(0.0) if total_upper > 0 else None
        current = 0
        
        for kp, count in selected_plan.items():
            if count <= 0:
                continue
            pool = [r for r in filtered if kp in (r.get("knowledge_points") or [])]
            if not pool:
                continue
            pool_with_time = [(r, r.get("created_time", 0)) for r in pool if r.get("handwriting_text") or r.get("attachments")]
            if not pool_with_time:
                continue
            pool_with_time.sort(key=lambda x: x[1], reverse=True)
            X = min(count, len(pool_with_time))
            reference_questions = [pool_with_time[i][0] for i in range(X)]
            
            generated_questions = []
            for ref in reference_questions:
                try:
                    texts = generate_similar_questions_with_llm(ref, 1, llm_api_key, llm_api_base, llm_model, token)
                    if texts:
                        generated_questions.append({
                            "subject": ref.get("subject"),
                            "knowledge_points": [kp],
                            "handwriting_text": texts[0],
                            "reason_type": "",
                            "reason_detail": "",
                            "attachments": [],
                            "created_time": 0,
                        })
                except Exception as e:
                    st.error(f"生成失败：{str(e)}")
                current += 1
                if progress_bar:
                    progress_bar.progress(min(1.0, current / total_upper))
            
            if generated_questions:
                similar_selections[kp] = generated_questions
        
        if progress_bar:
            progress_bar.progress(1.0)
        return similar_selections
    
    st.markdown("---")
    st.markdown("### 生成原题试卷")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("生成 Word 文档", type="primary", use_container_width=True, key="exam_word"):
            try:
                progress_bar = st.progress(0, text="正在准备题目...")
                selections = prepare_selections()
                if not selections:
                    st.warning("没有可用题目")
                    return
                progress_bar.progress(30, text="正在生成文档...")
                doc_bytes = build_doc(selected_subjects, selections, token)
                progress_bar.progress(100, text="生成完成！")
                filename = f"{'、'.join(selected_subjects)}_原题试卷.docx"
                st.success("✓ 生成成功")
                st.download_button("📥 下载 Word", data=doc_bytes, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            except Exception as e:
                st.error(f"生成失败：{e}")
    
    with col2:
        if st.button("生成 HTML 文档", type="primary", use_container_width=True, key="exam_html"):
            try:
                progress_bar = st.progress(0, text="正在准备题目...")
                selections = prepare_selections()
                if not selections:
                    st.warning("没有可用题目")
                    return
                progress_bar.progress(30, text="正在生成文档...")
                html_content = build_html(selected_subjects, selections, token)
                progress_bar.progress(100, text="生成完成！")
                filename = f"{'、'.join(selected_subjects)}_原题试卷.html"
                st.success("✓ 生成成功")
                st.download_button("📥 下载 HTML", data=html_content.encode('utf-8'), file_name=filename, mime="text/html", use_container_width=True)
            except Exception as e:
                st.error(f"生成失败：{e}")
    
    st.markdown("---")
    st.markdown("### 生成类似题试卷")
    
    if not llm_api_key:
        st.warning("⚠️ 需要配置智谱AI API Key 才能生成类似题目")
    else:
        col3, col4 = st.columns(2)
        with col3:
            if st.button("生成类似题 Word", type="primary", use_container_width=True, key="exam_similar_word"):
                try:
                    st.info("正在使用 AI 生成类似题目，请稍候...")
                    similar_selections = prepare_similar_selections()
                    if not similar_selections:
                        st.warning("生成失败或没有可用题目")
                        return
                    doc_bytes = build_doc(selected_subjects, similar_selections, token)
                    filename = f"{'、'.join(selected_subjects)}_类似题试卷.docx"
                    st.success("✓ 生成成功")
                    st.download_button("📥 下载 Word", data=doc_bytes, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="dl_similar_word")
                except Exception as e:
                    st.error(f"生成失败：{e}")
        
        with col4:
            if st.button("生成类似题 HTML", type="primary", use_container_width=True, key="exam_similar_html"):
                try:
                    st.info("正在使用 AI 生成类似题目，请稍候...")
                    similar_selections = prepare_similar_selections()
                    if not similar_selections:
                        st.warning("生成失败或没有可用题目")
                        return
                    html_content = build_html(selected_subjects, similar_selections, token)
                    filename = f"{'、'.join(selected_subjects)}_类似题试卷.html"
                    st.success("✓ 生成成功")
                    st.download_button("📥 下载 HTML", data=html_content.encode('utf-8'), file_name=filename, mime="text/html", use_container_width=True, key="dl_similar_html")
                except Exception as e:
                    st.error(f"生成失败：{e}")
    
    # 底部返回按钮
    st.markdown("---")
    if st.button("← 返回主页", key="exam_back_bottom"):
        st.session_state["current_page"] = "home"
        st.rerun()


def main() -> None:
    st.set_page_config(page_title="错题本", page_icon="📚", layout="wide")
    
    # 初始化页面状态
    if "current_page" not in st.session_state:
        st.session_state["current_page"] = "home"
    
    # 加载配置
    app_id, app_secret, llm_api_key, llm_api_base, llm_model, config, is_streamlit_cloud = _load_app_config()
    
    # 检查凭据
    app_id, app_secret = _check_feishu_credentials(app_id, app_secret, is_streamlit_cloud)
    
    # 主页不需要加载数据
    if st.session_state["current_page"] == "home":
        _render_home_page()
        return
    
    # 其他页面需要加载数据
    try:
        token = get_tenant_access_token(app_id, app_secret)
        raw_records = fetch_records(token)
        records = parse_records(raw_records)
    except requests.exceptions.ConnectionError as exc:
        st.error(f"网络连接失败：{exc}")
        if st.button("返回主页"):
            st.session_state["current_page"] = "home"
            st.rerun()
        return
    except RuntimeError as exc:
        msg = str(exc)
        if "99991663" in msg:
            st.error("飞书访问令牌无效，请检查应用权限配置")
        else:
            st.error(f"加载数据失败：{exc}")
        if st.button("返回主页"):
            st.session_state["current_page"] = "home"
            st.rerun()
        return
    except Exception as exc:
        st.error(f"加载数据失败：{exc}")
        if st.button("返回主页"):
            st.session_state["current_page"] = "home"
            st.rerun()
        return
    
    if not records:
        st.warning("表格暂无记录，请先在飞书多维表格填充数据。")
        if st.button("返回主页"):
            st.session_state["current_page"] = "home"
            st.rerun()
        return
    
    # 根据当前页面渲染内容
    if st.session_state["current_page"] == "practice":
        _render_practice_page(token, records, llm_api_key, llm_api_base, llm_model, config)
    elif st.session_state["current_page"] == "exam":
        _render_exam_page(token, records, llm_api_key, llm_api_base, llm_model)


if __name__ == "__main__":
    main()
